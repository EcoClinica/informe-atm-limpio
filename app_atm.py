import streamlit as st
import streamlit.components.v1 as components
from docxtpl import DocxTemplate
import datetime
import io
import re
import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- SISTEMA DE AUTO-GENERACIÓN DE PLANTILLA CORREGIDA (ARIAL Y AZUL CLARO) ---
def generar_plantilla_fiel():
    doc = Document()
    
    # Márgenes profesionales
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Color Azul Claro para títulos importantes
    AZUL_CLARO = RGBColor(2, 132, 199)
    NEGRO = RGBColor(0, 0, 0)
    GRIS_LINEA = RGBColor(156, 163, 175)
    
    # Título Principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INFORME ECOGRÁFICO DE LA ARTICULACIÓN TEMPOROMANDIBULAR\n(ATM)")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = NEGRO
    
    # Línea de separación inicial
    p_linea = doc.add_paragraph()
    r_l1 = p_linea.add_run("--------------------------------------------------------------------------------")
    r_l1.font.name = 'Arial'
    r_l1.font.color.rgb = GRIS_LINEA
    
    # Bloque Datos Paciente
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    
    r = p.add_run("Paciente: ")
    r.bold = True
    r.font.name = 'Arial'
    r.font.color.rgb = NEGRO
    
    r = p.add_run("{{ paciente }}                                                                             ")
    r.font.name = 'Arial'
    
    r = p.add_run("Edad: ")
    r.bold = True
    r.font.name = 'Arial'
    r.font.color.rgb = NEGRO
    
    r = p.add_run("{{ edad }}\n")
    r.font.name = 'Arial'
    
    r = p.add_run("Fecha: ")
    r.bold = True
    r.font.name = 'Arial'
    r.font.color.rgb = NEGRO
    
    r = p.add_run("{{ fecha }}                                                                     ")
    r.font.name = 'Arial'
    
    r = p.add_run("Derivado por: ")
    r.bold = True
    r.font.name = 'Arial'
    r.font.color.rgb = NEGRO
    
    r = p.add_run("{{ derivado }}\n")
    r.font.name = 'Arial'
    
    r = p.add_run("Motivo de consulta: ")
    r.bold = True
    r.font.name = 'Arial'
    r.font.color.rgb = NEGRO
    
    r = p.add_run("{{ motivo }}")
    r.font.name = 'Arial'
    
    # NUEVA LÍNEA DISCONTINUA DE SEPARACIÓN (Datos Paciente vs Estudio)
    p_linea_desc = doc.add_paragraph()
    p_linea_desc.paragraph_format.space_before = Pt(6)
    p_linea_desc.paragraph_format.space_after = Pt(12)
    r_ldesc = p_linea_desc.add_run("--------------------------------------------------------------------------------")
    r_ldesc.font.name = 'Arial'
    r_ldesc.font.color.rgb = GRIS_LINEA
    
    # Título general estudio (CENTRADO Y SIN REPETICIONES)
    p_estudio = doc.add_paragraph()
    p_estudio.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_estudio.paragraph_format.space_after = Pt(12)
    r = p_estudio.add_run("ESTUDIO DINÁMICO AMBAS ATM")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.font.color.rgb = NEGRO
    
    # --- FUNCIÓN REUTILIZABLE PARA AMBOS LADOS (Formato Arial y Azul Claro) ---
    def agregar_bloque_atm(lado_nombre, prefijo):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"ESTUDIO ARTICULACIÓN TEMPOROMANDIBULAR {lado_nombre}")
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.font.color.rgb = AZUL_CLARO # Aplicado azul claro solicitado
        
        # Campos de texto estructurados linealmente
        campos = [
            ("Cóndilo mandibular: ", f"{{{{ condilo_{prefijo} }}}}\n"),
            ("Espacio articular: ", f"{{{{ espacio_{prefijo} }}}}\n"),
            ("Derrame articular: ", f"{{{{ derrame_{prefijo} }}}}\n")
        ]
        
        p_campos = doc.add_paragraph()
        p_campos.paragraph_format.space_after = Pt(4)
        for label, var in campos:
            r1 = p_campos.add_run(label)
            r1.bold = True
            r1.font.name = 'Arial'
            r2 = p_campos.add_run(var)
            r2.font.name = 'Arial'
            
        # Fila de Medidas Condilares
        p_medidas = doc.add_paragraph()
        p_medidas.paragraph_format.space_after = Pt(4)
        r = p_medidas.add_run("Medidas condilares: ")
        r.bold = True
        r.font.name = 'Arial'
        
        r = p_medidas.add_run("Anterior ")
        r.italic = True
        r.font.name = 'Arial'
        r = p_medidas.add_run(f"{{{{ med_as_{prefijo} }}}} mm.   ")
        r.font.name = 'Arial'
        
        r = p_medidas.add_run("Lateral: ")
        r.italic = True
        r.font.name = 'Arial'
        r = p_medidas.add_run(f"{{{{ med_lat_{prefijo} }}}} mm.   ")
        r.font.name = 'Arial'
        
        r = p_medidas.add_run("Posterior: ")
        r.italic = True
        r.font.name = 'Arial'
        r = p_medidas.add_run(f"{{{{ med_pi_{prefijo} }}}} mm.\n")
        r.font.name = 'Arial'
        
        # Continuación de campos de posición
        r = p_medidas.add_run("Posición condilar (Pullinger): ")
        r.bold = True
        r.font.name = 'Arial'
        r = p_medidas.add_run(f"{{{{ pullinger_{prefijo} }}}}\n")
        r.font.name = 'Arial'
        
        r = p_medidas.add_run("Relación cóndilo-fosa glenoidea: ")
        r.bold = True
        r.font.name = 'Arial'
        r = p_medidas.add_run(f"{{{{ relacion_{prefijo} }}}}")
        r.font.name = 'Arial'
        
        # Disco articular y Hora en una sola línea espaciada
        p_disco = doc.add_paragraph()
        p_disco.paragraph_format.space_before = Pt(6)
        p_disco.paragraph_format.space_after = Pt(6)
        r = p_disco.add_run("Disco articular: ")
        r.bold = True
        r.font.name = 'Arial'
        r = p_disco.add_run(f"{{{{ disco_{prefijo} }}}}                                                           ")
        r.font.name = 'Arial'
        r = p_disco.add_run("Situación en hora: ")
        r.bold = True
        r.font.name = 'Arial'
        r = p_disco.add_run(f"{{{{ hora_{prefijo} }}}}")
        r.font.name = 'Arial'
        
        # Subsección de Estudio Dinámico
        p_dinamico = doc.add_paragraph()
        p_dinamico.paragraph_format.space_before = Pt(6)
        p_dinamico.paragraph_format.space_after = Pt(4)
        r = p_dinamico.add_run("Estudio dinámico:\n")
        r.bold = True
        r.font.name = 'Arial'
        
        r = p_dinamico.add_run("·       Boca cerrada: ")
        r.bold = True
        r.font.name = 'Arial'
        r = p_dinamico.add_run(f"{{{{ cerrada_{prefijo} }}}}\n")
        r.font.name = 'Arial'
        
        r = p_dinamico.add_run("·       Boca abierta: ")
        r.bold = True
        r.font.name = 'Arial'
        r = p_dinamico.add_run(f"{{{{ abierta_{prefijo} }}}}\n")
        r.font.name = 'Arial'
        
        r = p_dinamico.add_run("·       Reposición: ")
        r.bold = True
        r.font.name = 'Arial'
        r = p_dinamico.add_run(f"{{{{ repo_{prefijo} }}}}")
        r.font.name = 'Arial'
        
    agregar_bloque_atm("DERECHA", "der")
    agregar_bloque_atm("IZQUIERDA", "izq")
    
    # Línea final de separación
    p_linea2 = doc.add_paragraph()
    p_linea2.paragraph_format.space_before = Pt(12)
    r_l2 = p_linea2.add_run("--------------------------------------------------------------------------------")
    r_l2.font.name = 'Arial'
    r_l2.font.color.rgb = GRIS_LINEA
    
    # Conclusión en Azul Claro
    p_conclusion = doc.add_paragraph()
    r = p_conclusion.add_run("CONCLUSIÓN: ")
    r.bold = True
    r.font.name = 'Arial'
    r.font.color.rgb = AZUL_CLARO # Conclusión también resaltada en azul claro
    r = p_conclusion.add_run("{{ conclusion }}")
    r.font.name = 'Arial'
    
    bio_plantilla = io.BytesIO()
    doc.save(bio_plantilla)
    bio_plantilla.seek(0)
    return bio_plantilla

# Forzar la escritura física en el servidor
try:
    plantilla_binaria = generar_plantilla_fiel()
    with open("plantilla_atm.docx", "wb") as f:
        f.write(plantilla_binaria.getbuffer())
except Exception as e:
    pass

# Configuración de página ancha profesional para la web
st.set_page_config(page_title="Informe Ecográfico ATM", layout="wide", page_icon="🎙️")

# Estilo CSS para la interfaz web
st.markdown("""
    <style>
    .titulo-principal { color: #1E3A8A; font-weight: bold; text-align: center; margin-top: -20px; }
    .sub-seccion { color: #0284C7; border-bottom: 2px solid #0284C7; padding-bottom: 5px; margin-bottom: 15px; font-size: 20px; }
    .titulo-medidas { font-size: 14px; font-weight: bold; margin-bottom: 5px; color: #1E3A8A !important; }
    .esfera { font-size: 16px; vertical-align: middle; margin-right: 5px; }
    .resultado-calculo { background-color: #E0F2FE; padding: 12px; border-radius: 5px; border-left: 4px solid #0284C7; margin-top: 10px; margin-bottom: 15px; font-size: 14px; color: #1E3A8A !important; font-weight: bold; }
    </style>
""", unsafe_allow_html=True)

# --- MENÚ LATERAL: DESCARGAR PLANTILLA PARA RESPALDO ---
st.sidebar.markdown("### 📂 Zona de Respaldos")
plantilla_descarga = generar_plantilla_fiel()
st.sidebar.download_button(
    label="📥 Descargar plantilla original (.docx)",
    data=plantilla_descarga,
    file_name="plantilla_atm.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
st.sidebar.markdown("Use este botón si desea archivar este modelo limpio corregido en su iPad o computadora.")

st.markdown("<h1 class='titulo-principal'>Informe Ecográfico de la Articulación Temporomandibular (ATM)</h1>", unsafe_allow_html=True)
st.markdown("<hr style='border: 1px solid #1E3A8A;'>", unsafe_allow_html=True)

# Inicializar estados de dictado limpios
if "dictado_der" not in st.session_state: st.session_state.dictado_der = ""
if "dictado_izq" not in st.session_state: st.session_state.dictado_izq = ""

# --- COMPONENTE MICRÓFONO ---
def componente_microfono_visible(lado_id):
    js_code = f"""
    <div style="font-family: sans-serif; display: flex; flex-direction: column; gap: 5px;">
        <div style="display: flex; gap: 10px; align-items: center;">
            <button id="btn_{lado_id}" class="btn-voz btn-azul" onclick="conmutarMicro('{lado_id}')" style="flex-shrink: 0;">🎙️ Dictar 3 Medidas</button>
            <input type="text" id="output_local_{lado_id}" placeholder="Esperando dictado..." readonly 
                   style="flex-grow: 1; padding: 6px; font-size: 15px; font-weight: bold; border: 2px solid #0284C7; border-radius: 4px; background-color: #FFFFFF; color: #000000; text-align: center;">
        </div>
        <p id="status_{lado_id}" style="font-size:11px; color:#666; margin: 2px 0 0 2px; height: 14px; overflow: hidden;">Micro listo</p>
    </div>

    <script>
    let recognition_{lado_id} = null;
    let activo_{lado_id} = false;

    function enviarAStreamlit(textoNumeros) {{
        if (window.Streamlit) {{
            Streamlit.setComponentValue(textoNumeros);
        }}
    }}

    function conmutarMicro(lado) {{
        const btn = document.getElementById('btn_' + lado);
        const status = document.getElementById('status_' + lado);
        const inputLocal = document.getElementById('output_local_' + lado);

        if (activo_{lado_id}) {{
            if (recognition_{lado_id}) recognition_{lado_id}.abort();
            resetearBoton(lado, "🛑 Dictado cancelado.");
            return;
        }}

        if (!('webkitSpeechRecognition' in window) && !('SpeechRecognition' in window)) {{
            status.innerText = "Navegador no compatible.";
            return;
        }}

        const SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
        recognition_{lado_id} = new SpeechRecognition();
        recognition_{lado_id}.lang = 'es-ES';
        recognition_{lado_id}.interimResults = false;
        recognition_{lado_id}.maxAlternatives = 1;

        activo_{lado_id} = true;
        btn.innerText = "🛑 Cancelar";
        btn.className = "btn-voz btn-rojo";
        status.innerText = "🎙️ Escuchando... Di los 3 números seguidos";
        status.style.color = "#0284C7";
        inputLocal.value = "";

        recognition_{lado_id}.start();

        recognition_{lado_id}.onresult = function(event) {{
            const texto = event.results[0][0].transcript;
            const matches = texto.replace(/,/g, '.').match(/[0-9]+(\\.[0-9]+)?/g);
            
            if (matches && matches.length >= 3) {{
                const resultadoCadena = matches[0] + " , " + matches[1] + " , " + matches[2];
                inputLocal.value = resultadoCadena;
                status.innerText = "✓ Medidas capturadas correctamente.";
                status.style.color = "#16A34A";
                enviarAStreamlit(resultadoCadena);
            }} else {{
                status.innerText = "❌ Reintenta: Di 3 números claros.";
                status.style.color = "#DC2626";
            }}
        }};

        recognition_{lado_id}.onerror = function(e) {{
            if (e.error !== 'aborted') {{
                status.innerText = "❌ Error de micro o silencio.";
                status.style.color = "#DC2626";
            }}
            resetearBoton(lado, status.innerText);
        }};
        
        recognition_{lado_id}.onend = function() {{
            resetearBoton(lado, status.innerText);
        }};
    }}

    function resetearBoton(lado, msg) {{
        activo_{lado_id} = false;
        const btn = document.getElementById('btn_' + lado);
        const status = document.getElementById('status_' + lado);
        btn.innerText = "🎙️ Dictar 3 Medidas";
        btn.className = "btn-voz btn-azul";
        if(msg) status.innerText = msg;
    }}

    (function() {{
        var stScript = document.createElement('script');
        stScript.src = "https://cdn.jsdelivr.net/npm/@streamlit/component-lib@1.4.0/dist/index.min.js";
        stScript.onload = function() {{
            window.addEventListener('load', function() {{
                Streamlit.setFrameHeight(65);
            }});
        }};
        document.head.appendChild(stScript);
    }})();
    </script>

    <style>
    .btn-voz {{ border: none; padding: 8px 14px; border-radius: 4px; cursor: pointer; font-weight: bold; font-size: 13px; transition: background 0.2s; height: 35px; }}
    .btn-azul {{ background-color: #0284C7; color: white; }}
    .btn-azul:hover {{ background-color: #0369A1; }}
    .btn-rojo {{ background-color: #DC2626; color: white; }}
    .btn-rojo:hover {{ background-color: #B91C1C; }}
    </style>
    """
    return components.html(js_code, height=65, scrolling=False)

def procesar_medidas_sistema(texto_dictado, manual_as, manual_lat, manual_pi):
    if texto_dictado and isinstance(texto_dictado, str) and texto_dictado.strip():
        numeros = re.findall(r"[0-9]+(?:\.[0-9]+)?", texto_dictado)
        if len(numeros) >= 3:
            return numeros[0], numeros[1], numeros[2]
    return manual_as, manual_lat, manual_pi

def calcular_posicion_condilo(ant_sup_txt, post_inf_txt):
    try:
        if not ant_sup_txt or not post_inf_txt:
            return "Esperando medidas..."
        as_val = float(str(ant_sup_txt).replace(',', '.'))
        pi_val = float(str(post_inf_txt).replace(',', '.'))
        if (pi_val + as_val) == 0: return "0.00%"
        resultado = ((pi_val - as_val) / (pi_val + as_val)) * 100
        signo = "+" if resultado > 0 else ""
        return f"{signo}{resultado:.2f}%"
    except ValueError:
        return "Medidas pendientes"

# --- DATOS GENERALES ---
st.subheader("📋 Datos del Paciente")
with st.container(border=True):
    cp1, cp2, cp3 = st.columns(3)
    with cp1:
        nombres = st.text_input("Paciente (Nombre y Apellidos):")
        edad = st.text_input("Edad:")
    with cp2:
        fecha = st.date_input("Fecha:", datetime.date.today(), format="DD/MM/YYYY")
        derivado = st.text_input("Derivado por:")
    with cp3:
        motivo = st.text_input("Motivo de consulta:")

st.markdown("<br>", unsafe_allow_html=True)

# Opciones con primer elemento vacío para forzar inicio limpio
opts_condilo = ["", "Redondeado", "Aplanado", "En pico de pájaro (en punta)", "Con cresta central", "Con cresta marginal"]
opts_espacio = ["", "Libre", "Con engrosamiento sinovial", "Osteofitos", "Regular", "Irregular"]
opts_derrame = ["", "Sin derrame articular", "Con derrame anecoico", "Con derrame articular y con partículas ecogénicas"]
opts_relacion = ["", "Central", "Anterior", "Posterior"]
opts_disco = ["", "Homogénea, hipoecogénico", "Heterogénea", "Irregular"]
opts_horas = ["", "12", "11", "10", "1", "2"]
opts_boca_cerrada = ["", "Cubre totalmente la cabeza del cóndilo", "Cubre parcialmente la cabeza del cóndilo", "Desplazamiento, no cubre la cabeza condilar"]
opts_boca_abierta = [
    "",
    "Desplazamiento discal normal, cubre la cabeza del cóndilo",
    "Desplazamiento anterior, el disco cubre parcialmente la cabeza del cóndilo",
    "Desplazamiento anterior con subluxación discal, el cóndilo contacta con la cavidad glenoidea",
    "Desplazamiento anterior con luxación discal, el cóndilo contacta con la cavidad glenoidea",
    "Desplazamiento posterior, el disco cubre parcialmente la cabeza del cóndilo",
    "Desplazamiento posterior con subluxación discal, el cóndilo contacta con la cavidad glenoidea",
    "Desplazamiento posterior con luxación discal, el cóndilo contacta con la cavidad glenoidea"
]
opts_repo = ["", "Espontánea", "Requiere maniobras mandibulares para su recaptación por parte del paciente", "Requiere maniobras mandibulares para su recaptación por parte del médico", "No se reposiciona"]

# --- COLUMNAS PARALELAS EN LA WEB ---
col_der, col_izq = st.columns(2)

# --- ATM DERECHA ---
with col_der:
    with st.container(border=True):
        st.markdown("<h2 class='sub-seccion'><span class='esfera'>🔹</span>ATM Derecha</h2>", unsafe_allow_html=True)
        condilo_der = st.selectbox("Cóndilo mandibular (D):", opts_condilo, index=0, key="m_der")
        espacio_der = st.selectbox("Espacio articular (D):", opts_espacio, index=0, key="es_der")
        derrame_der = st.selectbox("Derrame articular (D):", opts_derrame, index=0, key="der_der")
        
        st.markdown("<p class='titulo-medidas'>Medidas condilares (mm):</p>", unsafe_allow_html=True)
        res_micro_der = componente_microfono_visible("der")
        if res_micro_der: st.session_state.dictado_der = res_micro_der
            
        m1, m2, m3 = st.columns(3)
        with m1: manual_as_der = st.text_input("Anterior (D)", value="", key="man_as_der")
        with m2: manual_lat_der = st.text_input("Lateral (D)", value="", key="man_lat_der")
        with m3: manual_pi_der = st.text_input("Posterior (D)", value="", key="man_pi_der")
        
        med_as_der, med_lat_der, med_pi_der = procesar_medidas_sistema(st.session_state.dictado_der, manual_as_der, manual_lat_der, manual_pi_der)
        res_der = calcular_posicion_condilo(med_as_der, med_pi_der)
        st.markdown(f"<div class='resultado-calculo'><strong>🧮 Posición condilar (Pullinger) (D):</strong> {res_der}</div>", unsafe_allow_html=True)
        
        relacion_der = st.selectbox("Relación cóndilo-fosa glenoidea (D):", opts_relacion, index=0, key="rel_der")
        disco_der = st.selectbox("Disco articular (D):", opts_disco, index=0, key="disco_der")
        hora_cerrada_der = st.selectbox("Situación en hora (D)", opts_horas, index=0, key="h_c_der")
        
        st.subheader("Estudio dinámico (D)")
        cerrada_txt_der = st.selectbox("Boca cerrada (D):", opts_boca_cerrada, index=0, key="c_txt_der")
        open_txt_der = st.selectbox("Boca abierta (D):", opts_boca_abierta, index=0, key="a_txt_der")
        repo_der = st.selectbox("Reposición (D):", opts_repo, index=0, key="rt_der")

# --- ATM IZQUIERDA ---
with col_izq:
    with st.container(border=True):
        st.markdown("<h2 class='sub-seccion'><span class='esfera'>🔹</span>ATM Izquierda</h2>", unsafe_allow_html=True)
        condilo_izq = st.selectbox("Cóndilo mandibular (I):", opts_condilo, index=0, key="m_izq")
        espacio_izq = st.selectbox("Espacio articular (I):", opts_espacio, index=0, key="es_izq")
        derrame_izq = st.selectbox("Derrame articular (I):", opts_derrame, index=0, key="der_izq")
        
        st.markdown("<p class='titulo-medidas'>Medidas condilares (mm):</p>", unsafe_allow_html=True)
        res_micro_izq = componente_microfono_visible("izq")
        if res_micro_izq: st.session_state.dictado_izq = res_micro_izq
            
        m4, m5, m6 = st.columns(3)
        with m4: manual_as_izq = st.text_input("Anterior (I)", value="", key="man_as_izq")
        with m5: manual_lat_izq = st.text_input("Lateral (I)", value="", key="man_lat_izq")
        with m6: manual_pi_izq = st.text_input("Posterior (I)", value="", key="man_pi_izq")
        
        med_as_izq, med_lat_izq, med_pi_izq = procesar_medidas_sistema(st.session_state.dictado_izq, manual_as_izq, manual_lat_izq, manual_pi_izq)
        res_izq = calcular_posicion_condilo(med_as_izq, med_pi_izq)
        st.markdown(f"<div class='resultado-calculo'><strong>🧮 Posición condilar (Pullinger) (I):</strong> {res_izq}</div>", unsafe_allow_html=True)
        
        relacion_izq = st.selectbox("Relación cóndilo-fosa glenoidea (I):", opts_relacion, index=0, key="rel_izq")
        disco_izq = st.selectbox("Disco articular (I):", opts_disco, index=0, key="disco_izq")
        hora_cerrada_izq = st.selectbox("Situación en hora (I)", opts_horas, index=0, key="h_c_izq")
        
        st.subheader("Estudio dinámico (I)")
        cerrada_txt_izq = st.selectbox("Boca cerrada (I):", opts_boca_cerrada, index=0, key="c_txt_izq")
        open_txt_izq = st.selectbox("Boca abierta (I):", opts_boca_abierta, index=0, key="a_txt_izq")
        repo_izq = st.selectbox("Reposición (I):", opts_repo, index=0, key="rt_izq")

st.markdown("<br><hr style='border: 1px solid #1E3A8A;'>", unsafe_allow_html=True)
conclusion = st.text_area("📝 CONCLUSIÓN:", height=120)

# --- EXPORTAR A WORD ---
try:
    doc = DocxTemplate("plantilla_atm.docx")
    contexto = {
        'paciente': nombres, 'edad': edad, 'derivado': derivado,
        'fecha': fecha.strftime("%d/%m/%Y"), 'motivo': motivo,
        
        'condilo_der': condilo_der, 'espacio_der': espacio_der, 'derrame_der': derrame_der,
        'med_as_der': med_as_der, 'med_lat_der': med_lat_der, 'med_pi_der': med_pi_der,
        'pullinger_der': res_der, 'relacion_der': relacion_der, 'disco_der': disco_der, 'hora_der': hora_cerrada_der,
        'cerrada_der': cerrada_txt_der, 'abierta_der': open_txt_der, 'repo_der': repo_der,
        
        'condilo_izq': condilo_izq, 'espacio_izq': espacio_izq, 'derrame_izq': derrame_izq,
        'med_as_izq': med_as_izq, 'med_lat_izq': med_lat_izq, 'med_pi_izq': med_pi_izq,
        'pullinger_izq': res_izq, 'relacion_izq': relacion_izq, 'disco_izq': disco_izq, 'hora_izq': hora_cerrada_izq,
        'cerrada_izq': cerrada_txt_izq, 'abierta_izq': open_txt_izq, 'repo_izq': repo_izq,
        
        'conclusion': conclusion
    }
    
    doc.render(contexto)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    st.download_button(
        label="🚀 DESCARGAR INFORME EN WORD",
        data=bio,
        file_name=f"Informe_ATM_{nombres}.docx".replace(" ", "_"),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
except Exception as e:
    st.error(f"Error al preparar el documento: {e}")
