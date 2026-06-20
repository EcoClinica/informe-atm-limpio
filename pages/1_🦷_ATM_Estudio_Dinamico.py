¡Entendido perfectamente! Vamos a pulir todos esos detalles para que el documento de Word quede impecable y la aplicación responda tal como lo necesitas en la práctica clínica.

Aquí tienes el análisis rápido de las correcciones aplicadas:

Desplegables multi-selección: Cambiados a st.multiselect para que puedas marcar varias opciones a la vez (se unirán con comas en el Word).

Limpieza por defecto: Si dejas un campo vacío en la app, en el Word saldrá completamente en blanco (ya no se arrastrarán los textos internos de la plantilla).

Formato Estricto Arial & Azul: Todo el documento de Word se genera en fuente Arial. Todos los títulos en mayúsculas cambian automáticamente a color azul, y el título principal queda centrado con un salto de línea estético y equilibrado.

📄 Copia todo este código corregido para pages/1_🦷_ATM_Estudio_Dinamico.py:
Python
import streamlit as st
import streamlit.components.v1 as components
from docxtpl import DocxTemplate
import datetime
import io
import re
import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_plantilla_fiel(ctx_datos):
    doc = Document()
    
    # Configuración de márgenes
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    AZUL_CLINICA = RGBColor(2, 132, 199)
    NEGRO = RGBColor(0, 0, 0)
    GRIS_LINEA = RGBColor(156, 163, 175)
    
    # 1. Título principal equilibrado en Arial y Azul
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(12)
    r_t1 = p_titulo.add_run("INFORME ECOGRÁFICO DE LA ARTICULACIÓN\nTEMPOROMANDIBULAR (ATM)")
    r_t1.bold = True
    r_t1.font.name = 'Arial'
    r_t1.font.size = Pt(14)
    r_t1.font.color.rgb = AZUL_CLINICA
    
    p_linea = doc.add_paragraph()
    r_l1 = p_linea.add_run("--------------------------------------------------------------------------------")
    r_l1.font.name = 'Arial'
    r_l1.font.color.rgb = GRIS_LINEA
    
    # 2. Datos del Paciente
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    
    def add_campo(parrafo, etiqueta, valor, espacio="   "):
        r_etiq = parrafo.add_run(etiqueta)
        r_etiq.bold = True
        r_etiq.font.name = 'Arial'
        r_val = parrafo.add_run(f"{valor}{espacio}")
        r_val.font.name = 'Arial'

    add_campo(p, "Paciente: ", ctx_datos.get('paciente', ''), "                                                                             ")
    add_campo(p, "Edad: ", ctx_datos.get('edad', ''), "\n")
    add_campo(p, "Fecha: ", ctx_datos.get('fecha', ''), "                                                                     ")
    add_campo(p, "Derivado por: ", ctx_datos.get('derivado', ''), "\n")
    add_campo(p, "Motivo de consulta: ", ctx_datos.get('motivo', ''), "")
    
    p_linea_desc = doc.add_paragraph()
    p_linea_desc.paragraph_format.space_before = Pt(6)
    p_linea_desc.paragraph_format.space_after = Pt(12)
    r_ldesc = p_linea_desc.add_run("--------------------------------------------------------------------------------")
    r_ldesc.font.name = 'Arial'
    r_ldesc.font.color.rgb = GRIS_LINEA
    
    # 3. Subtítulo General en Azul y Mayúsculas
    p_estudio = doc.add_paragraph()
    p_estudio.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_est = p_estudio.add_run("ESTUDIO DINÁMICO DE AMBAS ATM")
    r_est.bold = True
    r_est.font.name = 'Arial'
    r_est.font.size = Pt(12)
    r_est.font.color.rgb = AZUL_CLINICA
    
    # Bloques dinámicos de las ATM
    def agregar_bloque_atm(lado_nombre, prefijo):
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(12)
        r_sub = p_sub.add_run(f"ESTUDIO ARTICULACIÓN TEMPOROMANDIBULAR {lado_nombre}")
        r_sub.bold = True
        r_sub.font.name = 'Arial'
        r_sub.font.size = Pt(11)
        r_sub.font.color.rgb = AZUL_CLINICA
        
        p_campos = doc.add_paragraph()
        p_campos.paragraph_format.space_after = Pt(4)
        
        add_campo(p_campos, "Cóndilo mandibular: ", ctx_datos.get(f'condilo_{prefijo}', ''), "\n")
        add_campo(p_campos, "Espacio articular: ", ctx_datos.get(f'espacio_{prefijo}', ''), "\n")
        add_campo(p_campos, "Derrame articular: ", ctx_datos.get(f'derrame_{prefijo}', ''), "")
            
        p_med = doc.add_paragraph()
        p_med.paragraph_format.space_after = Pt(4)
        r_m_tit = p_med.add_run("Medidas condilares: ")
        r_m_tit.bold = True
        r_m_tit.font.name = 'Arial'
        
        r_ant = p_med.add_run("Anterior ")
        r_ant.italic = True
        r_ant.font.name = 'Arial'
        p_med.add_run(f"{ctx_datos.get(f'med_as_{prefijo}', '')} mm.  ").font.name = 'Arial'
        
        r_lat = p_med.add_run("Lateral: ")
        r_lat.italic = True
        r_lat.font.name = 'Arial'
        p_med.add_run(f"{ctx_datos.get(f'med_lat_{prefijo}', '')} mm.  ").font.name = 'Arial'
        
        r_post = p_med.add_run("Posterior: ")
        r_post.italic = True
        r_post.font.name = 'Arial'
        p_med.add_run(f"{ctx_datos.get(f'med_pi_{prefijo}', '')} mm.\n").font.name = 'Arial'
        
        add_campo(p_med, "Posición condilar (Pullinger): ", ctx_datos.get(f'pullinger_{prefijo}', ''), "\n")
        add_campo(p_med, "Relación cóndilo-fosa glenoidea: ", ctx_datos.get(f'relacion_{prefijo}', ''), "")
        
        p_disc = doc.add_paragraph()
        p_disc.paragraph_format.space_after = Pt(4)
        add_campo(p_disc, "Disco articular: ", ctx_datos.get(f'disco_{prefijo}', ''), "                                                           ")
        add_campo(p_disc, "Situación en hora: ", ctx_datos.get(f'hora_{prefijo}', ''), "")
        
        p_din = doc.add_paragraph()
        r_din_tit = p_din.add_run("Estudio dinámico:\n")
        r_din_tit.bold = True
        r_din_tit.font.name = 'Arial'
        
        add_campo(p_din, "·       Boca cerrada: ", ctx_datos.get(f'cerrada_{prefijo}', ''), "\n")
        add_campo(p_din, "·       Boca abierta: ", ctx_datos.get(f'abierta_{prefijo}', ''), "\n")
        add_campo(p_din, "·       Reposición: ", ctx_datos.get(f'repo_{prefijo}', ''), "")
        
    agregar_bloque_atm("DERECHA", "der")
    agregar_bloque_atm("IZQUIERDA", "izq")
    
    p_linea2 = doc.add_paragraph()
    p_linea2.add_run("--------------------------------------------------------------------------------").font.color.rgb = GRIS_LINEA
    
    p_c = doc.add_paragraph()
    r_c = p_c.add_run("CONCLUSIÓN: ")
    r_c.bold = True
    r_c.font.name = 'Arial'
    r_c.font.color.rgb = AZUL_CLINICA
    p_c.add_run(ctx_datos.get('conclusion', '')).font.name = 'Arial'
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

st.set_page_config(page_title="Informe Ecográfico ATM", layout="wide", page_icon="🎙️")

st.markdown("""
    <style>
    .titulo-principal { color: #1E3A8A; font-weight: bold; text-align: center; margin-top: -20px; }
    .sub-seccion { color: #0284C7; border-bottom: 2px solid #0284C7; padding-bottom: 5px; margin-bottom: 15px; font-size: 20px; }
    .titulo-medidas { font-size: 14px; font-weight: bold; margin-bottom: 5px; color: #1E3A8A !important; }
    .resultado-calculo { background-color: #E0F2FE; padding: 12px; border-radius: 5px; border-left: 4px solid #0284C7; margin-top: 10px; margin-bottom: 15px; font-size: 14px; color: #1E3A8A !important; font-weight: bold; }
    </style>
""", unsafe_allow_html=True)

st.markdown("<h1 class='titulo-principal'>Informe Ecográfico de la Articulación Temporomandibular (ATM)</h1>", unsafe_allow_html=True)

if "dictado_der" not in st.session_state: st.session_state.dictado_der = ""
if "dictado_izq" not in st.session_state: st.session_state.dictado_izq = ""

def componente_microfono_visible(lado_id):
    js_code = f"""
    <div style="font-family: sans-serif; display: flex; gap: 10px; align-items: center;">
        <button id="btn_{lado_id}" onclick="conmutarMicro('{lado_id}')" style="background-color:#0284C7; color:white; border:none; padding:6px 12px; border-radius:4px; font-weight:bold; cursor:pointer;">🎙️ Dictar 3 Medidas</button>
        <input type="text" id="out_{lado_id}" readonly style="padding:4px; border:2px solid #0284C7; border-radius:4px; text-align:center; width:150px;">
    </div>
    <script>
    let rec_{lado_id} = null; let act_{lado_id} = false;
    function conmutarMicro(l) {{
        const btn = document.getElementById('btn_'+l); const inp = document.getElementById('out_'+l);
        if(act_{lado_id}) {{ rec_{lado_id}.abort(); return; }}
        const SR = window.SpeechRecognition || window.webkitSpeechRecognition; if(!SR) return;
        rec_{lado_id} = new SR(); rec_{lado_id}.lang='es-ES'; act_{lado_id}=true; btn.innerText="🛑 Parar";
        rec_{lado_id}.start();
        rec_{lado_id}.onresult = function(e) {{
            const t = e.results[0][0].transcript.replace(/,/g, '.');
            const m = t.match(/[0-9]+(\\.[0-9]+)?/g);
            if(m && m.length >= 3) {{ inp.value = m[0]+" , "+m[1]+" , "+m[2]; if(window.Streamlit) Streamlit.setComponentValue(inp.value); }}
        }};
        rec_{lado_id}.onend = function() {{ act_{lado_id}=false; btn.innerText="🎙️ Dictar 3 Medidas"; }};
    }}
    (function() {{
        var s = document.createElement('script'); s.src = "https://cdn.jsdelivr.net/npm/@streamlit/component-lib@1.4.0/dist/index.min.js";
        s.onload = function() {{ window.addEventListener('load', function() {{ Streamlit.setFrameHeight(40); }}); }}; document.head.appendChild(s);
    }})();
    </script>
    """
    return components.html(js_code, height=40)

st.subheader("📋 Datos del Paciente")
with st.container(border=True):
    cp1, cp2, cp3 = st.columns(3)
    with cp1:
        nombres = st.text_input("Paciente:")
        edad = st.text_input("Edad:")
    with cp2:
        fecha = st.date_input("Fecha:", datetime.date.today(), format="DD/MM/YYYY")
        derivado = st.text_input("Derivado por:")
    with cp3:
        motivo = st.text_input("Motivo de consulta:")

opts_condilo = ["Redondeado", "Aplanado", "En pico de pájaro (en punta)", "Con cresta central", "Con cresta marginal"]
opts_espacio = ["Libre", "Con engrosamiento sinovial", "Osteofitos", "Regular", "Irregular"]
opts_derrame = ["Sin derrame articular", "Con derrame anecoico", "Con derrame articular y con partículas ecogénicas"]
opts_relacion = ["Central", "Anterior", "Posterior"]
opts_disco = ["Homogénea, hipoecogénico", "Heterogénea", "Irregular"]
opts_horas = ["12", "11", "10", "1", "2"]
opts_boca_cerrada = ["Cubre totalmente la cabeza del cóndilo", "Cubre parcialmente la cabeza del cóndilo", "Desplazamiento, no cubre la cabeza condilar"]
opts_boca_abierta = ["Desplazamiento discal normal, cubre la cabeza del cóndilo", "Desplazamiento anterior, el disco cubre parcialmente la cabeza del cóndilo", "Desplazamiento anterior con subluxación discal", "Desplazamiento anterior con luxación discal"]
opts_repo = ["Espontánea", "Requiere maniobras del paciente", "Requiere maniobras del médico", "No se reposiciona"]

col_der, col_izq = st.columns(2)

with col_der:
    with st.container(border=True):
        st.markdown("<h2 class='sub-seccion'>🔹ATM Derecha</h2>", unsafe_allow_html=True)
        condilo_der = st.multiselect("Cóndilo (D):", opts_condilo, key="c_d")
        espacio_der = st.multiselect("Espacio (D):", opts_espacio, key="e_d")
        derrame_der = st.multiselect("Derrame (D):", opts_derrame, key="d_d")
        
        st.markdown("<p class='titulo-medidas'>Medidas condilares (mm):</p>", unsafe_allow_html=True)
        r_md = componente_microfono_visible("der")
        if r_md: st.session_state.dictado_der = r_md
        
        m1, m2, m3 = st.columns(3)
        with m1: ma_d = st.text_input("Anterior (D)", key="ma_d")
        with m2: ml_d = st.text_input("Lateral (D)", key="ml_d")
        with m3: mp_d = st.text_input("Posterior (D)", key="mp_d")
        
        def proc(d, a, l, p):
            texto_dictado = str(d) if d is not None else ""
            if texto_dictado and len(re.findall(r"[0-9]+", texto_dictado)) >= 3: 
                return re.findall(r"[0-9.]+", texto_dictado)[:3]
            return a, l, p
            
        v1, v2, v3 = proc(st.session_state.dictado_der, ma_d, ml_d, mp_d)
        
        def calc(a, p):
            try:
                if not a or not p: return "Pendiente"
                v_a, v_p = float(str(a).replace(',','.')), float(str(p).replace(',','.'))
                return f"{'+' if (v_p-v_a)>0 else ''}{((v_p-v_a)/(v_p+v_a))*100:.2f}%"
            except: return "Error"
        res_d = calc(v1, v3)
        st.markdown(f"<div class='resultado-calculo'>🧮 Pullinger (D): {res_d}</div>", unsafe_allow_html=True)
        
        rel_d = st.multiselect("Relación (D):", opts_relacion, key="r_d")
        disc_d = st.multiselect("Disco (D):", opts_disco, key="di_d")
        h_d = st.multiselect("Hora (D):", opts_horas, key="h_d")
        cerr_d = st.multiselect("Boca cerrada (D):", opts_boca_cerrada, key="ce_d")
        ab_d = st.multiselect("Boca abierta (D):", opts_boca_abierta, key="ab_d")
        rep_d = st.multiselect("Reposición (D):", opts_repo, key="re_d")

with col_izq:
    with st.container(border=True):
        st.markdown("<h2 class='sub-seccion'>🔹ATM Izquierda</h2>", unsafe_allow_html=True)
        condilo_izq = st.multiselect("Cóndilo (I):", opts_condilo, key="c_i")
        espacio_izq = st.multiselect("Espacio (I):", opts_espacio, key="e_i")
        derrame_izq = st.multiselect("Derrame (I):", opts_derrame, key="d_i")
        
        st.markdown("<p class='titulo-medidas'>Medidas condilares (mm):</p>", unsafe_allow_html=True)
        r_mi = componente_microfono_visible("izq")
        if r_mi: st.session_state.dictado_izq = r_mi
        
        m4, m5, m6 = st.columns(3)
        with m4: ma_i = st.text_input("Anterior (I)", key="ma_i")
        with m5: ml_i = st.text_input("Lateral (I)", key="ml_i")
        with m6: mp_i = st.text_input("Posterior (I)", key="mp_i")
        v4, v5, v6 = proc(st.session_state.dictado_izq, ma_i, ml_i, mp_i)
        res_i = calc(v4, v6)
        st.markdown(f"<div class='resultado-calculo'>🧮 Pullinger (I): {res_i}</div>", unsafe_allow_html=True)
        
        rel_i = st.multiselect("Relación (I):", opts_relacion, key="r_i")
        disc_i = st.multiselect("Disco (I):", opts_disco, key="di_i")
        h_i = st.multiselect("Hora (I):", opts_horas, key="h_i")
        cerr_i = st.multiselect("Boca cerrada (I):", opts_boca_cerrada, key="ce_i")
        ab_i = st.multiselect("Boca abierta (I):", opts_boca_abierta, key="ab_i")
        rep_i = st.multiselect("Reposición (I):", opts_repo, key="re_i")

st.markdown("<br>", unsafe_allow_html=True)
conclusion = st.text_area("📝 CONCLUSIÓN:")

# Unión limpia de las opciones multiselección
def unir_opciones(lista):
    return ", ".join(lista) if lista else ""

ctx = {
    'paciente': nombres, 'edad': edad, 'derivado': derivado, 'fecha': fecha.strftime("%d/%m/%Y"), 'motivo': motivo,
    'condilo_der': unir_opciones(condilo_der), 'espacio_der': unir_opciones(espacio_der), 'derrame_der': unir_opciones(derrame_der), 
    'med_as_der': v1, 'med_lat_der': v2, 'med_pi_der': v3, 'pullinger_der': res_d if res_d != "Pendiente" else "", 
    'relacion_der': unir_opciones(rel_d), 'disco_der': unir_opciones(disc_d), 'hora_der': unir_opciones(h_d), 
    'cerrada_der': unir_opciones(cerr_d), 'abierta_der': unir_opciones(ab_d), 'repo_der': unir_opciones(rep_d),
    
    'condilo_izq': unir_opciones(condilo_izq), 'espacio_izq': unir_opciones(espacio_izq), 'derrame_izq': unir_opciones(derrame_izq), 
    'med_as_izq': v4, 'med_lat_izq': v5, 'med_pi_izq': v6, 'pullinger_izq': res_i if res_i != "Pendiente" else "", 
    'relacion_izq': unir_opciones(rel_i), 'disco_izq': unir_opciones(disc_i), 'hora_izq': unir_opciones(h_i), 
    'cerrada_izq': unir_opciones(cerr_i), 'abierta_izq': unir_opciones(ab_i), 'repo_izq': unir_opciones(rep_i),
    'conclusion': conclusion
}

# Botón de descarga dinámico
word_generado = generar_plantilla_fiel(ctx)
st.download_button(
    label="🚀 DESCARGAR INFORME EN WORD", 
    data=word_generado, 
    file_name=f"Informe_ATM_{nombres if nombres else 'Paciente'}.docx", 
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
