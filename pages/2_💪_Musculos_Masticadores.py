import streamlit as st
import streamlit.components.v1 as components
import datetime
import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_plantilla_musculos(ctx_datos):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    AZUL_CLINICA = RGBColor(2, 132, 199)
    NEGRO = RGBColor(0, 0, 0)
    GRIS_LINEA = RGBColor(156, 163, 175)
    
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(12)
    r_t1 = p_titulo.add_run("INFORME ECOGRÁFICO DE LOS MÚSCULOS MASTICADORES")
    r_t1.bold = True
    r_t1.font.name = 'Arial'
    r_t1.font.size = Pt(14)
    r_t1.font.color.rgb = AZUL_CLINICA
    
    p_linea = doc.add_paragraph()
    r_l1 = p_linea.add_run("--------------------------------------------------------------------------------")
    r_l1.font.name = 'Arial'
    r_l1.font.color.rgb = GRIS_LINEA
    
    # --- TABLA INVISIBLE PARA DATOS DEL PACIENTE ---
    tabla_datos = doc.add_table(rows=2, cols=2)
    tabla_datos.autofit = False
    tabla_datos.columns[0].width = Inches(4.5)
    tabla_datos.columns[1].width = Inches(2.4)
    
    def agregar_celda_campo(celda, etiqueta, valor):
        p_celda = celda.paragraphs[0]
        p_celda.paragraph_format.space_after = Pt(4)
        r_etiq = p_celda.add_run(etiqueta)
        r_etiq.bold = True
        r_etiq.font.name = 'Arial'
        r_val = p_celda.add_run(str(valor))
        r_val.font.name = 'Arial'

    agregar_celda_campo(tabla_datos.cell(0, 0), "Paciente: ", ctx_datos.get('paciente', ''))
    agregar_celda_campo(tabla_datos.cell(0, 1), "Edad: ", ctx_datos.get('edad', ''))
    agregar_celda_campo(tabla_datos.cell(1, 0), "Fecha: ", ctx_datos.get('fecha', ''))
    agregar_celda_campo(tabla_datos.cell(1, 1), "Derivado por: ", ctx_datos.get('derivado', ''))
    
    p_motivo = doc.add_paragraph()
    p_motivo.paragraph_format.space_before = Pt(4)
    p_motivo.paragraph_format.space_after = Pt(8)
    r_mot_etiq = p_motivo.add_run("Motivo de consulta: ")
    r_mot_etiq.bold = True
    r_mot_etiq.font.name = 'Arial'
    r_mot_val = p_motivo.add_run(ctx_datos.get('motivo', ''))
    r_mot_val.font.name = 'Arial'
    
    p_linea_desc = doc.add_paragraph()
    p_linea_desc.paragraph_format.space_before = Pt(6)
    p_linea_desc.paragraph_format.space_after = Pt(12)
    r_ldesc = p_linea_desc.add_run("--------------------------------------------------------------------------------")
    r_ldesc.font.name = 'Arial'
    r_ldesc.font.color.rgb = GRIS_LINEA
    
    p_estudio = doc.add_paragraph()
    p_estudio.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_est = p_estudio.add_run("ESTUDIO ECOGRÁFICO MUSCULAR COMPARATIVO")
    r_est.bold = True
    r_est.font.name = 'Arial'
    r_est.font.size = Pt(12)
    r_est.font.color.rgb = AZUL_CLINICA
    
    def agregar_bloque_musculo(nombre_musculo, prefijo, ref_txt):
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(12)
        r_sub = p_sub.add_run(f"MÚSCULO {nombre_musculo.upper()}")
        r_sub.bold = True
        r_sub.font.name = 'Arial'
        r_sub.font.size = Pt(11)
        r_sub.font.color.rgb = AZUL_CLINICA
        
        p_campos = doc.add_paragraph()
        p_campos.paragraph_format.space_after = Pt(4)
        
        # 1. VALORES DE REFERENCIA DEL MÚSCULO EN EL WORD
        r_ref_tit = p_campos.add_run(f"Valores Ref. ({ref_txt}) | Fuente: Ultrasonography in TMD\n")
        r_ref_tit.font.name = 'Arial'
        r_ref_tit.font.size = Pt(8.5)
        r_ref_tit.italic = True
        r_ref_tit.font.color.rgb = RGBColor(107, 114, 128)
        
        # 2. MEDIDAS
        r_tit_med = p_campos.add_run("Espesor y Dinámica Muscular:\n")
        r_tit_med.bold = True
        r_tit_med.font.name = 'Arial'
        
        r_d_tit = p_campos.add_run("· Lado Derecho: ")
        r_d_tit.bold = True
        r_d_tit.font.name = 'Arial'
        r_d_val = p_campos.add_run(f"Reposo: {ctx_datos.get(f'rep_d_{prefijo}', '')} mm | Contracción: {ctx_datos.get(f'con_d_{prefijo}', '')} mm | Engrosamiento: {ctx_datos.get(f'pct_d_{prefijo}', '')}\n")
        r_d_val.font.name = 'Arial'
        
        r_i_tit = p_campos.add_run("· Lado Izquierdo: ")
        r_i_tit.bold = True
        r_i_tit.font.name = 'Arial'
        r_i_val = p_campos.add_run(f"Reposo: {ctx_datos.get(f'rep_i_{prefijo}', '')} mm | Contracción: {ctx_datos.get(f'con_i_{prefijo}', '')} mm | Engrosamiento: {ctx_datos.get(f'pct_i_{prefijo}', '')}\n\n")
        r_i_val.font.name = 'Arial'
        
        # 3. ECOESTRUCTURA, SIMETRÍA Y HALLAZGOS DESPUÉS
        def add_campo_linea(parrafo, etiqueta, valor):
            r_etiq = parrafo.add_run(etiqueta)
            r_etiq.bold = True
            r_etiq.font.name = 'Arial'
            r_val = parrafo.add_run(f"{valor}\n")
            r_val.font.name = 'Arial'
            
        add_campo_linea(p_campos, "Ecoestructura general: ", ctx_datos.get(f'eco_{prefijo}', ''))
        add_campo_linea(p_campos, "Simetría comparativa: ", ctx_datos.get(f'simetria_{prefijo}', ''))
        add_campo_linea(p_campos, "Hallazgos / Fasciculaciones: ", ctx_datos.get(f'hallazgos_{prefijo}', ''))

    agregar_bloque_musculo("Masetero", "mas", "Mujeres: R 8-10mm, C 10-12mm | Varones: R 10-12mm, C 13-15mm")
    agregar_bloque_musculo("Temporal", "tem", "Mujeres: R 3.8-4.8mm, C 5-6mm | Varones: R 4.5-5.5mm, C 6-7mm")
    
    p_linea2 = doc.add_paragraph()
    p_linea2.paragraph_format.space_before = Pt(12)
    p_linea2.add_run("--------------------------------------------------------------------------------").font.color.rgb = GRIS_LINEA
    
    p_c = doc.add_paragraph()
    r_c = p_c.add_run("CONCLUSIÓN GENERAL DEL ESTUDIO MUSCULAR: ")
    r_c.bold = True
    r_c.font.name = 'Arial'
    r_c.font.color.rgb = AZUL_CLINICA
    p_c.add_run(ctx_datos.get('conclusion', '')).font.name = 'Arial'
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

st.set_page_config(page_title="Ecografía Músculos Masticadores", layout="wide", page_icon="💪")

st.markdown("""
    <style>
    .titulo-principal { color: #1E3A8A; font-weight: bold; text-align: center; margin-top: -20px; }
    .sub-seccion { color: #0284C7; border-bottom: 2px solid #0284C7; padding-bottom: 5px; margin-bottom: 15px; font-size: 20px; }
    .titulo-medidas { font-size: 14px; font-weight: bold; margin-bottom: 5px; color: #1E3A8A !important; }
    .resultado-calculo { background-color: #F0FDF4; padding: 10px; border-radius: 5px; border-left: 4px solid #22C55E; margin-top: 5px; margin-bottom: 10px; font-size: 13px; color: #166534 !important; font-weight: bold; }
    .micro-referencia { color: #475569; font-size: 12px; margin-top: -8px; margin-bottom: 12px; font-style: italic; background-color: #F8FAFC; padding: 6px; border-radius: 4px; border: 1px solid #E2E8F0;}
    </style>
""", unsafe_allow_html=True)

st.markdown("<h1 class='titulo-principal'>Estudio Ecográfico de los Músculos Masticadores</h1>", unsafe_allow_html=True)

st.subheader("📋 Datos del Paciente")
with st.container(border=True):
    cp1, cp2, cp3, cp4 = st.columns([2, 1, 1, 2])
    with cp1:
        nombres = st.text_input("Paciente:")
    with cp2:
        edad = st.text_input("Edad:")
    with cp3:
        sexo = st.selectbox("Sexo:", ["Seleccionar", "Femenino", "Masculino"])
    with cp4:
        fecha = st.date_input("Fecha:", datetime.date.today(), format="DD/MM/YYYY")
        derivado = st.text_input("Derivado por:")

motivo = st.text_input("Motivo de consulta:")

opts_ecoestructura = ["Normal, patrón fibrilar conservado", "Hiperecoicidad difusa (Fibrosis/Sobrecarga)", "Pérdida del patrón fibrilar", "Zonas de atrofia muscular"]
opts_simetria = ["Simétrico bilateralmente", "Asimetría por hipertrofia derecha", "Asimetría por hipertrofia izquierda", "Atrofia unilateral"]
opts_hallazgos = ["Sin hallazgos patológicos relevantes", "Presencia de bandas tensas microrroturas", "Fasciculaciones evidente en dinámica", "Sialolitiasis en glándula parótida adyacente"]

col_mas, col_tem = st.columns(2)

def calcular_engrosamiento(reposo, contraccion):
    r = str(reposo).strip()
    c = str(contraccion).strip()
    if not r or not c:
        return "Pendiente"
    try:
        v_r = float(r.replace(',', '.'))
        v_c = float(c.replace(',', '.'))
        if v_r <= 0: return "Pendiente"
        resultado = ((v_c - v_r) / v_r) * 100
        return f"+{resultado:.2f}%" if resultado > 0 else f"{resultado:.2f}%"
    except ValueError:
        return "Pendiente"

# --- BLOQUE MÚSCULO MASETERO ---
with col_mas:
    with st.container(border=True):
        st.markdown("<h2 class='sub-seccion'>💪 Músculo Masetero</h2>", unsafe_allow_html=True)
        
        # Referencias abreviadas aquí dentro
        st.markdown("""
        <div class='micro-referencia'>
            <strong>📊 Valores Ref. Normales:</strong><br>
            • 🚺 <strong>Mujeres:</strong> R: 8.0-10.0 mm | C: 10.0-12.0 mm<br>
            • 🚹 <strong>Varones:</strong> R: 10.0-12.0 mm | C: 13.0-15.0 mm<br>
            <span style='font-size: 10px; color: #94A3B8;'>Fuente: Manfredini et al., Ultrasonography in TMD Diagnosis.</span>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<p class='titulo-medidas'>Espesor Lado Derecho (mm):</p>", unsafe_allow_html=True)
        md1, md2 = st.columns(2)
        with md1: rep_d_mas = st.text_input("Reposo (D)", value="", key="rd_mas")
        with md2: con_d_mas = st.text_input("Máx. Contracción (D)", value="", key="cd_mas")
        pct_d_mas = calcular_engrosamiento(rep_d_mas, con_d_mas)
        st.markdown(f"<div class='resultado-calculo'>📈 Engrosamiento Masetero (D): {pct_d_mas}</div>", unsafe_allow_html=True)
        
        st.markdown("<p class='titulo-medidas'>Espesor Lado Izquierdo (mm):</p>", unsafe_allow_html=True)
        mi1, mi2 = st.columns(2)
        with mi1: rep_i_mas = st.text_input("Reposo (I)", value="", key="ri_mas")
        with mi2: con_i_mas = st.text_input("Máx. Contracción (I)", value="", key="ci_mas")
        pct_i_mas = calcular_engrosamiento(rep_i_mas, con_i_mas)
        st.markdown(f"<div class='resultado-calculo'>📈 Engrosamiento Masetero (I): {pct_i_mas}</div>", unsafe_allow_html=True)
        
        eco_mas = st.multiselect("Ecoestructura (Masetero):", opts_ecoestructura, key="eco_m")
        sim_mas = st.multiselect("Simetría (Masetero):", opts_simetria, key="sim_m")
        hal_mas = st.multiselect("Hallazgos (Masetero):", opts_hallazgos, key="hal_m")

# --- BLOQUE MÚSCULO TEMPORAL ---
with col_tem:
    with st.container(border=True):
        st.markdown("<h2 class='sub-seccion'>💪 Músculo Temporal</h2>", unsafe_allow_html=True)
        
        # Referencias abreviadas aquí dentro
        st.markdown("""
        <div class='micro-referencia'>
            <strong>📊 Valores Ref. Normales:</strong><br>
            • 🚺 <strong>Mujeres:</strong> R: 3.8-4.8 mm | C: 5.0-6.0 mm<br>
            • 🚹 <strong>Varones:</strong> R: 4.5-5.5 mm | C: 6.0-7.0 mm<br>
            <span style='font-size: 10px; color: #94A3B8;'>Fuente: Diagnostic Criteria for TMD (DC/TMD) / Ultrasonography protocols.</span>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<p class='titulo-medidas'>Espesor Lado Derecho (mm):</p>", unsafe_allow_html=True)
        td1, td2 = st.columns(2)
        with td1: rep_d_tem = st.text_input("Reposo (D)", value="", key="rd_tem")
        with td2: con_d_tem = st.text_input("Máx. Contracción (D)", value="", key="cd_tem")
        pct_d_tem = calcular_engrosamiento(rep_d_tem, con_d_tem)
        st.markdown(f"<div class='resultado-calculo'>📈 Engrosamiento Temporal (D): {pct_d_tem}</div>", unsafe_allow_html=True)
        
        st.markdown("<p class='titulo-medidas'>Espesor Lado Izquierdo (mm):</p>", unsafe_allow_html=True)
        ti1, ti2 = st.columns(2)
        with ti1: rep_i_tem = st.text_input("Reposo (I)", value="", key="ri_tem")
        with ti2: con_i_tem = st.text_input("Máx. Contracción (I)", value="", key="ci_tem")
        pct_i_tem = calcular_engrosamiento(rep_i_tem, con_i_tem)
        st.markdown(f"<div class='resultado-calculo'>📈 Engrosamiento Temporal (I): {pct_i_tem}</div>", unsafe_allow_html=True)
        
        eco_tem = st.multiselect("Ecoestructura (Temporal):", opts_ecoestructura, key="eco_t")
        sim_tem = st.multiselect("Simetría (Temporal):", opts_simetria, key="sim_t")
        hal_tem = st.multiselect("Hallazgos (Temporal):", opts_hallazgos, key="hal_t")

st.markdown("<br>", unsafe_allow_html=True)
conclusion = st.text_area("📝 CONCLUSIÓN GENERAL:")

def unir_opciones(lista):
    return ", ".join(lista) if lista else ""

ctx = {
    'paciente': nombres, 'edad': edad, 'derivado': derivado, 'fecha': fecha.strftime("%d/%m/%Y"), 'motivo': motivo,
    'sexo': sexo if sexo != "Seleccionar" else "No especificado",
    
    # Masetero
    'eco_mas': unir_opciones(eco_mas), 'simetria_mas': unir_opciones(sim_mas), 'hallazgos_mas': unir_opciones(hal_mas),
    'rep_d_mas': rep_d_mas.strip(), 'con_d_mas': con_d_mas.strip(), 'pct_d_mas': pct_d_mas if pct_d_mas != "Pendiente" else "",
    'rep_i_mas': rep_i_mas.strip(), 'con_i_mas': con_i_mas.strip(), 'pct_i_mas': pct_i_mas if pct_i_mas != "Pendiente" else "",
    
    # Temporal
    'eco_tem': unir_opciones(eco_tem), 'simetria_tem': unir_opciones(sim_tem), 'hallazgos_tem': unir_opciones(hal_tem),
    'rep_d_tem': rep_d_tem.strip(), 'con_d_tem': con_d_tem.strip(), 'pct_d_tem': pct_d_tem if pct_d_tem != "Pendiente" else "",
    'rep_i_tem': rep_i_tem.strip(), 'con_i_tem': con_i_tem.strip(), 'pct_i_tem': pct_i_tem if pct_i_tem != "Pendiente" else "",
    
    'conclusion': conclusion
}

word_generado = generar_plantilla_musculos(ctx)
st.download_button(
    label="🚀 DESCARGAR INFORME MUSCULAR EN WORD", 
    data=word_generado, 
    file_name=f"Informe_Muscular_{nombres if nombres else 'Paciente'}.docx", 
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
