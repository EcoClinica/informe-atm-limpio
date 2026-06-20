import streamlit as st
import streamlit.components.v1 as components
from docxtpl import DocxTemplate
import datetime
import io
import re
import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_plantilla_musculos():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    AZUL_CLARO = RGBColor(2, 132, 199)
    NEGRO = RGBColor(0, 0, 0)
    GRIS_LINEA = RGBColor(156, 163, 175)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESTUDIO ECOGRÁFICO DE LOS MÚSCULOS MASTICADORES")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    
    p_linea = doc.add_paragraph()
    p_linea.add_run("--------------------------------------------------------------------------------").font.color.rgb = GRIS_LINEA
    
    p = doc.add_paragraph()
    p.add_run("Paciente: ").bold = True
    p.add_run("{{ paciente }}                                                                             ").font.name = 'Arial'
    p.add_run("Edad: ").bold = True
    p.add_run("{{ edad }}\n").font.name = 'Arial'
    p.add_run("Fecha: ").bold = True
    p.add_run("{{ fecha }}                                                                     ").font.name = 'Arial'
    p.add_run("Derivado por: ").bold = True
    p.add_run("{{ derivado }}\n").font.name = 'Arial'
    p.add_run("Motivo de consulta: ").bold = True
    p.add_run("{{ motivo }}").font.name = 'Arial'
    
    p_linea_desc = doc.add_paragraph()
    p_linea_desc.add_run("--------------------------------------------------------------------------------").font.color.rgb = GRIS_LINEA
    
    def agregar_bloque_musculo(titulo, prefijo, es_m=False):
        p = doc.add_paragraph()
        r = p.add_run(titulo)
        r.bold = True
        r.font.color.rgb = AZUL_CLARO
        r.font.size = Pt(11)
        
        p_c = doc.add_paragraph()
        p_c.add_run("Ecoestructura: ").bold = True
        p_c.add_run(f"{{{{ eco_{prefijo} }}}}\n").font.name = 'Arial'
        p_c.add_run("Músculo relajado diám AP: ").bold = True
        p_c.add_run(f"{{{{ med_rel_{prefijo} }}}} mm")
        if es_m: p_c.add_run(" (VN H: 10-15 mm) (VN M: 9 - 13 mm)")
        p_c.add_run("\n")
        p_c.add_run("Músculo en contracción máxima diám AP: ").bold = True
        p_c.add_run(f"{{{{ med_cont_{prefijo} }}}} mm")
        if es_m: p_c.add_run(" (VN H: 14-19 mm) (VN M: 12-15 mm)")
        
    agregar_bloque_musculo("Masetero derecho:", "mas_der", es_m=True)
    agregar_bloque_musculo("Temporal derecho:", "temp_der")
    agregar_bloque_musculo("Masetero izquierdo:", "mas_izq", es_m=True)
    agregar_bloque_musculo("Temporal izquierdo:", "temp_izq")
    
    p_l2 = doc.add_paragraph()
    p_l2.add_run("--------------------------------------------------------------------------------").font.color.rgb = GRIS_LINEA
    
    p_con = doc.add_paragraph()
    r = p_con.add_run("CONCLUSIÓN: ")
    r.bold = True
    r.font.color.rgb = AZUL_CLARO
    p_con.add_run("{{ conclusion }}")
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

try:
    with open("plantilla_musculos.docx", "wb") as f: f.write(generar_plantilla_musculos().getbuffer())
except: pass

st.set_page_config(page_title="Informe Músculos Masticadores", layout="wide", page_icon="💪")

st.sidebar.markdown("### 📂 Zona de Respaldos")
st.sidebar.download_button("📥 Descargar plantilla original (.docx)", generar_plantilla_musculos(), "plantilla_musculos.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.markdown("<h1 style='color: #1E3A8A; font-weight: bold; text-align: center;'>Estudio Ecográfico de los Músculos Masticadores</h1>", unsafe_allow_html=True)

st.subheader("📋 Datos del Paciente")
with st.container(border=True):
    cp1, cp2, cp3 = st.columns(3)
    with cp1:
        nombres = st.text_input("Paciente:")
        edad = st.text_input("Edad:")
    with cp2:
        fecha = st.date_input("Fecha:", datetime.date.today(), format="DD/MM/YYYY")
        derivado = st.text_input("Derivado por:")
    with cp3:
        motivo = st.text_input("Motivo de consulta:")
        genero = st.radio("Género:", ["Mujer", "Hombre"], horizontal=True)

opts_eco = ["Conservada", "Alterada", "Hipertrófico", "Atrófico"]
col_der, col_izq = st.columns(2)

with col_der:
    with st.container(border=True):
        st.subheader("🔹Hemicara Derecha")
        eco_md = st.selectbox("Ecoestructura (Masetero D):", opts_eco)
        m_r_md = st.text_input("Relajado (mm) [Masetero D]")
        m_c_md = st.text_input("Contracción (mm) [Masetero D]")
        
        st.markdown("---")
        eco_td = st.selectbox("Ecoestructura (Temporal D):", opts_eco)
        t_r_td = st.text_input("Relajado (mm) [Temporal D]")
        t_c_td = st.text_input("Contracción (mm) [Temporal D]")

with col_izq:
    with st.container(border=True):
        st.subheader("🔹Hemicara Izquierda")
        eco_mi = st.selectbox("Ecoestructura (Masetero I):", opts_eco)
        m_r_mi = st.text_input("Relajado (mm) [Masetero I]")
        m_c_mi = st.text_input("Contracción (mm) [Masetero I]")
        
        st.markdown("---")
        eco_ti = st.selectbox("Ecoestructura (Temporal I):", opts_eco)
        t_r_ti = st.text_input("Relajado (mm) [Temporal I]")
        t_c_ti = st.text_input("Contracción (mm) [Temporal I]")

conclusion = st.text_area("📝 CONCLUSIÓN:", value="Ecoestructura muscular conservada globalmente sin evidencias de asimetrías.")

try:
    doc = DocxTemplate("plantilla_musculos.docx")
    ctx = {
        'paciente': nombres, 'edad': edad, 'derivado': derivado, 'fecha': fecha.strftime("%d/%m/%Y"), 'motivo': motivo,
        'eco_mas_der': eco_md, 'med_rel_mas_der': m_r_md if m_r_md else ".....", 'med_cont_mas_der': m_c_md if m_c_md else ".....",
        'eco_temp_der': eco_td, 'med_rel_temp_der': t_r_td if t_r_td else ".....", 'med_cont_temp_der': t_c_td if t_c_td else ".....",
        'eco_mas_izq': eco_mi, 'med_rel_mas_izq': m_r_mi if m_r_mi else ".....", 'med_cont_mas_izq': m_c_mi if m_c_mi else ".....",
        'eco_temp_izq': eco_ti, 'med_rel_temp_izq': t_r_ti if t_r_ti else ".....", 'med_cont_temp_izq': t_c_ti if t_c_ti else ".....",
        'conclusion': conclusion
    }
    doc.render(ctx)
    bio_o = io.BytesIO()
    doc.save(bio_o)
    bio_o.seek(0)
    st.download_button("🚀 DESCARGAR INFORME MÚSCULOS EN WORD", bio_o, f"Informe_Musculos_{nombres}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
except Exception as e:
    st.error(f"Error: {e}")
